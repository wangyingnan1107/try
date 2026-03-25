"""
Word文档导出模块
将HTML内容转换为Word文档，保留格式和图片

Java vs Python 说明：
- python-docx库用于创建Word文档，类似Java的Apache POI
- BytesIO用于在内存中处理字节数据，类似Java的ByteArrayOutputStream
"""
import io
from typing import Optional
from bs4 import BeautifulSoup
import requests
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import re
import logging

logger = logging.getLogger(__name__)


class WordExporter:
    """
    Word文档导出器类
    类似Java: public class WordExporter { }
    """
    
    def __init__(self):
        """
        初始化导出器
        """
        self.headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
        }
    
    def download_image(self, url: str, base_url: str = None) -> Optional[bytes]:
        """
        下载图片
        
        Args:
            url: 图片URL（可能是相对路径或绝对路径）
            base_url: 基础URL，用于拼接相对路径
            
        Returns:
            Optional[bytes]: 图片的字节数据，失败返回None
        """
        try:
            # 如果是相对路径，使用base_url拼接
            # urljoin用于拼接URL，类似Java的URI.resolve()
            if base_url and not url.startswith(('http://', 'https://')):
                url = requests.compat.urljoin(base_url, url)
            
            # 发送HTTP请求下载图片
            response = requests.get(url, headers=self.headers, timeout=10, stream=True)
            response.raise_for_status()
            
            # 检查Content-Type是否为图片
            content_type = response.headers.get('Content-Type', '')
            if not content_type.startswith('image/'):
                logger.warning(f"URL不是图片: {url}, Content-Type: {content_type}")
                return None
            
            # 返回图片的字节数据
            return response.content
        except Exception as e:
            logger.error(f"下载图片失败 {url}: {e}")
            return None
    
    def html_to_word(self, html_content: str, base_url: str = None) -> bytes:
        """
        将HTML内容转换为Word文档
        
        Args:
            html_content: HTML内容（字符串）
            base_url: 基础URL，用于下载图片
            
        Returns:
            bytes: Word文档的字节数据
        """
        # 创建Word文档对象
        # 类似Java: XWPFDocument document = new XWPFDocument();
        doc = Document()
        
        # 设置中文字体
        # 需要设置中文字体，否则中文可能显示为方块
        self._set_chinese_font(doc)
        
        try:
            # 使用BeautifulSoup解析HTML
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # 移除不需要的标签
            for element in soup(['script', 'style', 'meta', 'link', 'noscript']):
                element.decompose()
            
            # 提取body内容，如果没有body则使用整个HTML
            body = soup.find('body') or soup
            
            # 遍历HTML元素并转换为Word格式
            self._process_element(body, doc, base_url)
            
        except Exception as e:
            logger.error(f"HTML转Word失败: {e}")
            # 如果转换失败，至少添加原始文本
            doc.add_paragraph(html_content)
        
        # 将Word文档保存到内存中的字节流
        # BytesIO用于在内存中处理字节数据，类似Java的ByteArrayOutputStream
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)  # 将文件指针移到开头
        
        return buffer.getvalue()
    
    def _set_chinese_font(self, doc: Document):
        """
        设置Word文档的中文字体
        
        Args:
            doc: Word文档对象
        """
        # 获取文档的样式
        # 类似Java: XWPFParagraph paragraph = document.createParagraph();
        style = doc.styles['Normal']
        
        # 设置字体
        # 类似Java: paragraph.setFontFamily("微软雅黑");
        font = style.font
        font.name = '微软雅黑'  # 中文字体
        font.size = Pt(12)  # 字体大小12磅
        
        # 设置中文字体（重要：否则中文可能显示异常）
        # qn用于处理命名空间，类似Java的XML命名空间处理
        font._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    
    def _process_element(self, element, doc: Document, base_url: str = None):
        """
        递归处理HTML元素，转换为Word格式
        
        Args:
            element: BeautifulSoup元素对象
            doc: Word文档对象
            base_url: 基础URL
        """
        # 遍历元素的直接子节点（不包括嵌套的子节点）
        # element.children只包含直接子节点
        for child in element.children:
            # 检查节点类型
            if hasattr(child, 'name'):  # 如果是标签节点
                tag_name = child.name.lower()
                
                # 处理标题标签
                if tag_name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                    self._add_heading(child, doc, tag_name)
                
                # 处理段落标签
                elif tag_name == 'p':
                    self._add_paragraph(child, doc, base_url)
                
                # 处理图片标签
                elif tag_name == 'img':
                    self._add_image(child, doc, base_url)
                
                # 处理列表
                elif tag_name in ['ul', 'ol']:
                    self._add_list(child, doc, base_url, tag_name == 'ol')
                
                # 处理换行
                elif tag_name == 'br':
                    doc.add_paragraph()  # 添加空段落作为换行
                
                # 处理div等其他块级元素（递归处理其内容）
                elif tag_name in ['div', 'article', 'section', 'main']:
                    self._process_element(child, doc, base_url)
                
                # 处理其他标签（提取文本内容）
                else:
                    text = child.get_text(strip=True)
                    if text:
                        para = doc.add_paragraph(text)
                        self._format_paragraph(para, child)
            
            # 如果是文本节点
            elif child.string and child.string.strip():
                # 获取文本内容
                text = child.string.strip()
                if text:
                    doc.add_paragraph(text)
    
    def _add_heading(self, element, doc: Document, level_tag: str):
        """
        添加标题
        
        Args:
            element: HTML元素
            doc: Word文档对象
            level_tag: 标题级别标签（h1-h6）
        """
        text = element.get_text(strip=True)
        if text:
            # 根据标签确定标题级别（h1=1级，h2=2级，以此类推）
            level = int(level_tag[1])  # 提取数字，如'h1'提取1
            doc.add_heading(text, level=min(level, 9))  # Word最多支持9级标题
    
    def _add_paragraph(self, element, doc: Document, base_url: str = None):
        """
        添加段落
        
        Args:
            element: HTML段落元素
            doc: Word文档对象
            base_url: 基础URL
        """
        # 获取段落中的所有内容（包括文本和子标签）
        para_text = []
        images_in_para = []
        
        for child in element.children:
            if hasattr(child, 'name'):
                if child.name == 'img':
                    images_in_para.append(child)
                else:
                    text = child.get_text(strip=True)
                    if text:
                        para_text.append(text)
            elif child.string and child.string.strip():
                para_text.append(child.string.strip())
        
        # 如果有文本，添加段落
        if para_text:
            para = doc.add_paragraph(' '.join(para_text))
            self._format_paragraph(para, element)
        
        # 如果有图片，添加图片
        for img_elem in images_in_para:
            self._add_image(img_elem, doc, base_url)
    
    def _add_image(self, img_element, doc: Document, base_url: str = None):
        """
        添加图片到Word文档
        
        Args:
            img_element: HTML img元素
            doc: Word文档对象
            base_url: 基础URL
        """
        # 获取图片URL（src属性）
        img_url = img_element.get('src') or img_element.get('data-src')
        
        if not img_url:
            return
        
        try:
            # 下载图片
            img_data = self.download_image(img_url, base_url)
            
            if img_data:
                # 将图片添加到文档中
                # BytesIO用于将字节数据转换为文件对象
                img_stream = io.BytesIO(img_data)
                
                # 添加到段落中
                para = doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 居中对齐
                
                # 添加图片到段落（最大宽度6英寸）
                run = para.add_run()
                run.add_picture(img_stream, width=Inches(6))
            
        except Exception as e:
            logger.error(f"添加图片失败 {img_url}: {e}")
            # 如果添加图片失败，至少添加图片的alt文本
            alt_text = img_element.get('alt', '')
            if alt_text:
                doc.add_paragraph(f"[图片: {alt_text}]")
    
    def _add_list(self, list_element, doc: Document, base_url: str = None, ordered: bool = False):
        """
        添加列表
        
        Args:
            list_element: HTML列表元素（ul或ol）
            doc: Word文档对象
            base_url: 基础URL
            ordered: 是否为有序列表（True=有序，False=无序）
        """
        # 查找所有li元素
        items = list_element.find_all('li', recursive=False)
        
        for item in items:
            text = item.get_text(strip=True)
            if text:
                # 添加段落作为列表项
                para = doc.add_paragraph(text, style='List Bullet' if not ordered else 'List Number')
                self._format_paragraph(para, item)
                
                # 检查列表项中是否有图片
                for img in item.find_all('img'):
                    self._add_image(img, doc, base_url)
    
    def _format_paragraph(self, para, element):
        """
        根据HTML元素的样式格式化段落
        
        Args:
            para: Word段落对象
            element: HTML元素
        """
        # 获取样式信息
        style = element.get('style', '')
        class_name = element.get('class', [])
        
        # 如果是字符串，转换为列表
        if isinstance(class_name, str):
            class_name = [class_name]
        
        # 处理粗体、斜体等格式（需要检查父元素或子元素）
        # 这里简化处理，实际可以更复杂
        if element.find(['strong', 'b']):
            for run in para.runs:
                run.bold = True
        
        if element.find(['em', 'i']):
            for run in para.runs:
                run.italic = True
        
        # 处理对齐方式
        if 'text-align: center' in style or 'center' in class_name:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif 'text-align: right' in style or 'right' in class_name:
            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    def text_to_word(self, text: str, title: str = "提取的文本") -> bytes:
        """
        将纯文本转换为Word文档（简单版本，不包含格式）
        
        Args:
            text: 文本内容
            title: 文档标题
            
        Returns:
            bytes: Word文档的字节数据
        """
        doc = Document()
        
        # 设置中文字体
        self._set_chinese_font(doc)
        
        # 添加标题
        doc.add_heading(title, 0)
        
        # 按段落分割文本
        paragraphs = text.split('\n\n')
        
        for para_text in paragraphs:
            if para_text.strip():
                # 按行分割，处理换行
                lines = para_text.split('\n')
                for line in lines:
                    if line.strip():
                        doc.add_paragraph(line.strip())
                # 段落之间添加空行
                doc.add_paragraph()
        
        # 保存到内存
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        return buffer.getvalue()


def html_to_word_doc(html_content: str, base_url: str = None) -> bytes:
    """
    便捷函数：将HTML转换为Word文档
    
    Args:
        html_content: HTML内容
        base_url: 基础URL
        
    Returns:
        bytes: Word文档的字节数据
    """
    exporter = WordExporter()
    return exporter.html_to_word(html_content, base_url)


def text_to_word_doc(text: str, title: str = "提取的文本") -> bytes:
    """
    便捷函数：将文本转换为Word文档
    
    Args:
        text: 文本内容
        title: 文档标题
        
    Returns:
        bytes: Word文档的字节数据
    """
    exporter = WordExporter()
    return exporter.text_to_word(text, title)

